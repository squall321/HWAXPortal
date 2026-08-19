# 대화 이력을 Word(.docx) 로 만든다 — 전문(있는 그대로)과 정리본(LLM 재구성) 두 모드.
"""왜 서버에서 만드는가.

브라우저에서 .docx 를 만들려면 프론트에 새 의존성이 붙는다. python-docx 는 이미 사내
다른 앱이 쓰고 포털 백엔드에도 들어 있어, 서식 규칙을 한 곳에 모을 수 있다.

두 모드가 성격이 다르다.
  transcript — 있는 그대로. 누가 무엇을 말했는지가 기록으로 남아야 하는 경우.
  report     — LLM 이 읽고 재구성. 결론이 먼저 오고 근거가 따라오는 문서.

report 는 LLM 왕복이 있어 느리고, LLM 이 죽으면 실패한다. 그때 조용히 transcript 로
떨어지지 않는다 — 정리본을 기대한 사람에게 원문을 주면 '정리가 안 된 문서'로 보이지
'실패'로 보이지 않기 때문이다.
"""
from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# 본문 색 — 인쇄를 전제로 한다. 화면용 밝은 회색은 종이에서 사라진다.
_INK = RGBColor(0x1A, 0x1A, 0x1A)
_MUTED = RGBColor(0x66, 0x66, 0x66)
_RULE = "─" * 46


def _style(doc: Document) -> None:
    """기본 서식 — 본문 10.5pt, 줄간격 넉넉히. 한글 폰트는 문서 기본에 맡긴다."""
    n = doc.styles["Normal"]
    n.font.size = Pt(10.5)
    n.font.color.rgb = _INK
    pf = n.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.35


def _meta(doc: Document, title: str, subtitle: str) -> None:
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = _INK

    s = doc.add_paragraph()
    sr = s.add_run(subtitle)
    sr.font.size = Pt(9)
    sr.font.color.rgb = _MUTED

    d = doc.add_paragraph()
    dr = d.add_run(_RULE)
    dr.font.color.rgb = _MUTED
    d.paragraph_format.space_after = Pt(14)


def _md_para(doc: Document, line: str) -> None:
    """마크다운 한 줄을 문단으로. 제목·목록·인용만 다룬다 — 표는 원문 그대로 둔다."""
    t = line.rstrip()
    if not t.strip():
        return
    m = re.match(r"^(#{1,4})\s+(.*)$", t)
    if m:
        doc.add_heading(m.group(2).strip(), level=min(len(m.group(1)), 4))
        return
    if re.match(r"^\s*[-*•]\s+", t):
        doc.add_paragraph(re.sub(r"^\s*[-*•]\s+", "", t), style="List Bullet")
        return
    m = re.match(r"^\s*(\d+)[.)]\s+(.*)$", t)
    if m:
        doc.add_paragraph(m.group(2), style="List Number")
        return
    if t.lstrip().startswith(">"):
        p = doc.add_paragraph(t.lstrip().lstrip(">").strip())
        p.paragraph_format.left_indent = Pt(18)
        for r in p.runs:
            r.italic = True
            r.font.color.rgb = _MUTED
        return
    p = doc.add_paragraph()
    # **굵게** 만 인라인으로 살린다. 나머지 기호는 그대로 둬야 원문이 훼손되지 않는다.
    for i, seg in enumerate(re.split(r"\*\*(.+?)\*\*", t)):
        if not seg:
            continue
        run = p.add_run(seg)
        run.bold = i % 2 == 1


# XML(따라서 docx)이 담을 수 없는 제어문자. 대화 본문에 ANSI 이스케이프나 NUL 이 한 글자만
# 있어도 python-docx 가 예외를 던져 내보내기 전체가 500 이 된다 — 로그를 붙여넣은 대화에서
# 실제로 흔하다. 탭·개행은 남기고 나머지만 지운다.
_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _clean(text: str) -> str:
    return _CTRL.sub("", text or "")


def _body(doc: Document, text: str) -> None:
    for line in _clean(text).splitlines():
        _md_para(doc, line)


def _ts(ms: int | None) -> str:
    if not ms:
        return ""
    return datetime.fromtimestamp(ms / 1000).strftime("%Y-%m-%d %H:%M")


def _delib_turns(m: dict[str, Any]) -> list[dict]:
    """심의 좌석 발언. 이것이 없으면 심의를 내보내도 결정문만 남는다.

    ⚠ 좌석 발언은 message.text 에 없다 — 라이브 스트림은 delib 이벤트로만 오고(m.delib.turns),
    서버에서 다시 읽을 때도 role='persona' 행이 turns 로 접혀 들어간다. text 만 읽으면
    3라운드 5좌석 심의가 '질문 + 결정문' 두 줄로 줄어든다. 결론만 남고 근거가 사라지는
    문서는 심의 기록으로 쓸 수 없다.
    """
    d = m.get("delib")
    if not isinstance(d, dict):
        return []
    # 원소까지 검사한다. _write_turns 는 걸렀는데 transcript_text 는 안 걸러서,
    # delib.turns 에 문자열이 하나만 섞여도 AttributeError 가 났다 — DocxRequest 가
    # extra=allow 라 임의 JSON 이 그대로 들어온다. 여기서 한 번만 거른다.
    return [t for t in (d.get("turns") or []) if isinstance(t, dict)]


def _write_turns(doc: "Document", turns: list[dict]) -> None:
    last_round = None
    for t in turns:
        if not isinstance(t, dict):
            continue
        rd = t.get("round")
        if rd != last_round:
            last_round = rd
            h = doc.add_paragraph()
            hr = h.add_run(_clean(f"라운드 {rd}") if rd is not None else "심의")
            hr.bold = True
            hr.font.size = Pt(10)
            hr.font.color.rgb = _MUTED
            h.paragraph_format.space_before = Pt(10)
        # 여기도 정화한다 — _body 만 걸어 두면 persona·stance·position 이 그대로
        # add_run 에 들어가 같은 ValueError 로 500 이 된다(실측).
        who = _clean(str(t.get("persona") or "")).strip() or "좌석"
        stance = _clean(str(t.get("stance") or "")).strip()
        pos = _clean(str(t.get("position") or "")).strip()
        lead = doc.add_paragraph()
        lr = lead.add_run(who + (f"  ·  {stance}" if stance else ""))
        lr.bold = True
        lr.font.size = Pt(9)
        if pos:
            pr = lead.add_run(f"   {pos}")
            pr.font.size = Pt(9)
            pr.font.color.rgb = _MUTED
        _body(doc, str(t.get("say") or ""))
        nn = _clean(str(t.get("nonNegotiable") or "")).strip()
        if nn:
            _body(doc, f"양보 불가: {nn}")


def build_transcript(conv: dict[str, Any]) -> bytes:
    """있는 그대로 — 발화자·시각을 붙여 순서대로."""
    doc = Document()
    _style(doc)
    # text 가 비어도 심의 발언이 있으면 버리지 않는다 — 결정문 없이 끝난 심의가 통째로
    # 사라지던 경로다.
    msgs = [m for m in (conv.get("messages") or [])
            if (m.get("text") or "").strip() or _delib_turns(m)]
    _meta(doc, _clean(conv.get("title") or "대화 기록"),
          f"HWAX Portal · 발화 {len(msgs)}개 · 내보낸 시각 "
          f"{datetime.now().strftime('%Y-%m-%d %H:%M')}")

    for m in msgs:
        who = {"user": "질문", "assistant": "응답"}.get(m.get("role"), m.get("role") or "")
        p = doc.add_paragraph()
        r = p.add_run(f"{who}   {_ts(m.get('ts'))}")
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = _MUTED
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(2)
        # 좌석 발언을 먼저 싣고 결정문(text)을 뒤에 둔다 — 읽는 순서가 심의 순서와 같아진다.
        _write_turns(doc, _delib_turns(m))
        _body(doc, m.get("text") or "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_report(title: str, markdown: str, source_note: str) -> bytes:
    title, markdown, source_note = _clean(title), _clean(markdown), _clean(source_note)
    """정리본 — LLM 이 만든 마크다운을 문서 서식으로 옮긴다."""
    doc = Document()
    _style(doc)
    _meta(doc, title or "정리 보고서",
          f"HWAX Portal · {source_note} · 작성 "
          f"{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    _body(doc, markdown)

    # 출처를 문서 안에 남긴다 — 이 문서가 대화에서 나왔다는 사실이 파일만 봐서는 안 보인다.
    doc.add_paragraph()
    p = doc.add_paragraph()
    r = p.add_run("※ 이 문서는 HWAX Portal 대화 이력을 정리해 자동 생성했습니다. "
                  "수치·인용은 원문에서 확인하십시오.")
    r.font.size = Pt(8.5)
    r.font.color.rgb = _MUTED

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def transcript_text(conv: dict[str, Any], limit: int = 60000) -> str:
    """LLM 에 넘길 대화 원문. 너무 길면 앞부분을 자르되 잘랐다는 사실을 남긴다."""
    out = []
    for m in conv.get("messages") or []:
        # 좌석 발언도 넣는다. 빼면 LLM 이 근거 없이 결론만 보고 '결론과 근거' 절을 쓰게 된다 —
        # 정리본이 지어낸 근거를 달게 되는 가장 나쁜 경로다.
        for t_ in _delib_turns(m):
            say = str(t_.get("say") or "").strip()
            if say:
                out.append(f"[{t_.get('persona') or '좌석'} R{t_.get('round')}] {say}")
        t = (m.get("text") or "").strip()
        if not t:
            continue
        who = {"user": "질문", "assistant": "응답"}.get(m.get("role"), m.get("role") or "")
        out.append(f"[{who}] {t}")
    full = "\n\n".join(out)
    if len(full) <= limit:
        return full
    return full[:limit] + "\n\n(…이후 생략 — 원문이 길어 앞부분만 전달했다)"
